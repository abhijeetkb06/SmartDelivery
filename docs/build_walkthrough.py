"""Build walkthrough Word document from captured screenshots."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS = Path(__file__).parent
OUTPUT = DOCS / "SmartDelivery_Walkthrough.docx"

SECTIONS = [
    {
        "title": "1. My myQ — Homeowner View",
        "image": "walkthrough_01_myq_home.png",
        "text": (
            "The My myQ tab is the homeowner-facing view. It displays the garage door status, "
            "the latest delivery intelligence alert, and a before-vs-after comparison showing "
            "the difference between raw myQ sensor events (left) and Couchbase-powered smart "
            "delivery alerts (right).\n\n"
            "Key features:\n"
            "- Real-time garage door status (Open/Closed)\n"
            "- Delivery Intelligence banner with scenario-aware messaging\n"
            "- Before vs After comparison: raw events vs AI-enriched alerts\n"
            "- Risk-aware color coding: red for at-risk deliveries, green for safe"
        ),
    },
    {
        "title": "2. My myQ — Active Alerts & Recent Activity",
        "image": "walkthrough_02_myq_alerts.png",
        "text": (
            "Scrolling down on the My myQ tab reveals active alerts and recent delivery activity. "
            "Each notification card shows a scenario-aware smart summary generated automatically "
            "by Couchbase Eventing — no application code needed.\n\n"
            "Key features:\n"
            "- Active Alerts with risk scores and actionable recommendations\n"
            "- Recent Delivery Activity cards with per-delivery smart summaries\n"
            "- 7 scenario types: happy path, package behind car, front door misdelivery, "
            "door stuck open, no package placed, delivery timeout, theft/suspicious\n"
            "- Expandable 'View Couchbase Eventing Pipeline' showing the 3-step enrichment process"
        ),
    },
    {
        "title": "3. myQ Command Center — Fleet Overview",
        "image": "walkthrough_03_command_center.png",
        "text": (
            "The Command Center is the fleet operations view. It shows aggregate statistics "
            "across all monitored homes and provides controls for the event generator.\n\n"
            "Key features:\n"
            "- 1,000,000 Homes Monitored, Total Deliveries, Alerts Triggered, AI-Ready Records\n"
            "- Start/Stop Event Stream button launching the Go event generator at 5,000 ops/sec\n"
            "- Auto-stop after 1 million deliveries per session\n"
            "- Live pipeline performance metrics with real-time ops/sec counter"
        ),
    },
    {
        "title": "4. Command Center — Pipeline Performance & PII Redaction",
        "image": "walkthrough_04_pipeline_pii.png",
        "text": (
            "The pipeline performance section shows the real-time data flow from ingestion "
            "through Couchbase Eventing to AI-ready records. Below it, the PII Redaction "
            "showcase demonstrates automatic name masking by Couchbase Eventing.\n\n"
            "Key features:\n"
            "- Live pipeline flow: Ingest -> Eventing -> Vector Embedding -> AI-Ready\n"
            "- Automatic PII redaction: 'John Smith' -> 'J*** S***' (server-side, zero app code)\n"
            "- Raw data vs Processed data side-by-side comparison\n"
            "- Addresses preserved for operations; only names are redacted"
        ),
    },
    {
        "title": "5. Command Center — Alert Feed",
        "image": "walkthrough_05_alert_feed.png",
        "text": (
            "The Alert Feed provides a filterable view of all delivery alerts across the fleet. "
            "Operators can filter by severity level to triage incidents.\n\n"
            "Key features:\n"
            "- Severity filter: All, Critical, High, Medium, Low\n"
            "- Color-coded alert cards with severity indicators\n"
            "- Alert types: theft risk, door stuck, package at risk, misdelivery, etc.\n"
            "- Address and timestamp for each alert"
        ),
    },
    {
        "title": "6. Vector Search & AI Copilot",
        "image": "walkthrough_06_vector_search.png",
        "text": (
            "The Vector Search & AI Copilot tab combines semantic delivery search with a "
            "RAG-powered conversational AI assistant.\n\n"
            "Delivery Search features:\n"
            "- Natural language queries (e.g., 'suspicious activity at garage door')\n"
            "- Filters: Carrier, Scenario, Status, Risk Level, Result Limit\n"
            "- Couchbase Vector Search with APPROX_VECTOR_DISTANCE (1,536-dim embeddings)\n"
            "- Expandable 'View Couchbase Query' showing the actual SQL++ vector query\n\n"
            "AI Copilot features:\n"
            "- Conversational RAG Q&A over delivery data\n"
            "- Streaming LLM responses (GPT-4o-mini) with cited delivery IDs\n"
            "- 5-step RAG pipeline: Embed Query -> Vector Search -> Build Context -> Feed to LLM -> Stream Response\n"
            "- Cached embeddings for repeat questions (1-hour TTL)"
        ),
    },
]


def main():
    doc = Document()

    # ── Title page ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("myQ Smart Delivery Intelligence", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Application Walkthrough\n"
        "AI-powered package delivery monitoring for your garage — powered by Couchbase"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # ── Architecture overview ──
    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph(
        "SmartDelivery demonstrates Couchbase's full-stack capabilities in a single application:\n\n"
        "1. Go Event Generator — Streams 5,000+ delivery events/sec using collection.Do() "
        "bulk API with 50 parallel goroutines\n"
        "2. Couchbase Eventing — Server-side functions automatically enrich raw deliveries: "
        "PII redaction, knowledge narrative generation, risk assessment\n"
        "3. Vector Embedding Pipeline — Second eventing function calls OpenAI to generate "
        "1,536-dim embeddings for semantic search\n"
        "4. Streamlit Dashboard — Three views (Homeowner, Operations, AI Search) powered by "
        "N1QL queries with GSI indexes, KV lookups, and vector search\n"
        "5. RAG AI Copilot — Conversational Q&A using Couchbase Vector Search for context "
        "retrieval and GPT-4o-mini for response generation"
    )

    doc.add_page_break()

    # ── Walkthrough sections ──
    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)

        img_path = DOCS / section["image"]
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer
        doc.add_paragraph(section["text"])
        doc.add_page_break()

    # ── Tech stack summary ──
    doc.add_heading("Technology Stack", level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "Component"
    headers[1].text = "Technology"
    data = [
        ("Database", "Couchbase Capella (Cloud)"),
        ("Event Generator", "Go 1.21+ with gocb/v2 SDK, 50 goroutines, collection.Do() bulk API"),
        ("Server-Side Logic", "Couchbase Eventing (JavaScript, automatic on mutation)"),
        ("Embeddings", "OpenAI text-embedding-3-small (1,536 dimensions)"),
        ("Vector Search", "Couchbase Vector Index with APPROX_VECTOR_DISTANCE"),
        ("LLM", "OpenAI GPT-4o-mini (streaming, RAG)"),
        ("Frontend", "Streamlit 1.53+ with streamlit-shadcn-ui components"),
    ]
    for i, (comp, tech) in enumerate(data):
        row = table.rows[i + 1].cells
        row[0].text = comp
        row[1].text = tech

    doc.save(str(OUTPUT))
    print(f"Walkthrough saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
